import io
import docx
import pdfplumber

def parse_docx(file_bytes: bytes) -> list:
    """
    Parses a DOCX file and extracts paragraphs, headers, and tables,
    preserving structural hierarchy and numbering.
    """
    doc = docx.Document(io.BytesIO(file_bytes))
    elements = []
    
    # Process paragraphs
    for idx, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue
            
        style = p.style.name if p.style else "Normal"
        is_heading = style.startswith("Heading") or (len(text) < 120 and (
            text.lower().startswith("section") or 
            text.lower().startswith("article") or 
            text.lower().startswith("clause") or
            (text[0].isdigit() if text else False)
        ) and not text.endswith(","))
        
        elements.append({
            "index": len(elements),
            "text": text,
            "type": "heading" if is_heading else "paragraph",
            "page": 1  # DOCX doesn't have page numbers naturally easily
        })
        
    # Process tables
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            # De-duplicate cell texts (merged cells might duplicate text)
            cells = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if not cells or cells[-1] != cell_text:
                    cells.append(cell_text)
            
            row_text = " | ".join(c for c in cells if c)
            if row_text:
                elements.append({
                    "index": len(elements),
                    "text": f"[Table {t_idx+1}, Row {r_idx+1}]: {row_text}",
                    "type": "table",
                    "page": 1
                })
                
    return elements

def parse_pdf(file_bytes: bytes) -> list:
    """
    Parses a PDF file using pdfplumber, preserving structural hierarchy, 
    numbering, and page numbers.
    """
    elements = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for p_idx, page in enumerate(pdf.pages):
            text = page.extract_text()
            if not text:
                continue
                
            lines = text.split("\n")
            for line in lines:
                line_str = line.strip()
                if not line_str:
                    continue
                
                # Check for headings
                is_heading = len(line_str) < 120 and (
                    line_str.lower().startswith("section") or 
                    line_str.lower().startswith("article") or 
                    line_str.lower().startswith("clause") or 
                    line_str.upper() == line_str or
                    (line_str[0].isdigit() if line_str else False)
                )
                
                elements.append({
                    "index": len(elements),
                    "text": line_str,
                    "type": "heading" if is_heading else "paragraph",
                    "page": p_idx + 1
                })
                
    return elements

def extract_document_data(file_bytes: bytes, filename: str) -> dict:
    """
    Extracts text, preserves structures, and detects if the document is scanned.
    """
    ext = filename.split(".")[-1].lower()
    elements = []
    is_scanned = False
    
    if ext == "docx":
        elements = parse_docx(file_bytes)
    elif ext == "pdf":
        elements = parse_pdf(file_bytes)
        # If no text is extracted from a PDF, it's likely scanned
        if not elements:
            is_scanned = True
    else:
        raise ValueError(f"Unsupported file format: {ext}")
        
    full_text = "\n".join(el["text"] for el in elements)
    
    # If text is extremely short for a document, flag as scanned/noisy
    if len(full_text.strip()) < 100 and ext == "pdf":
        is_scanned = True
        
    return {
        "filename": filename,
        "is_scanned": is_scanned,
        "elements": elements,
        "full_text": full_text
    }
